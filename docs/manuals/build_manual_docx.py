from __future__ import annotations

import csv
import re
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "manuals" / "dist"
IMG_DIR = OUT_DIR / "diagram_png"
DOCX_PATH = OUT_DIR / "DotTalkPP_x64base_Manual_Preview_V1.docx"

MANUAL = ROOT / "DOTTALKPP_READER_MANUAL_V1.md"
ANCHOR_MAP = ROOT / "DOTTALKPP_MANUAL_ANCHOR_MAP_V1.md"
MATURITY = ROOT / "DOTTALKPP_MANUAL_MATURITY_MODEL_V1.md"
COMMAND_GUIDE = ROOT / "DOTTALKPP_COMMAND_REFERENCE_GUIDE_V1.md"
DIAGRAM_REGISTRY = ROOT / "docs" / "manuals" / "assets" / "diagrams" / "manual_diagram_asset_registry_v1.csv"

ACCENT = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 96, 104)
LIGHT_FILL = "E8EEF5"
PALE_FILL = "F4F6F9"
BORDER = "B7C6D8"


DIAGRAMS = {
    "DIAG-ARCH-LAYERS": [
        ["DotTalk++ shell", "DotScript", "HELP/META", "Manuals"],
        ["x64base runtime", "DBF records", "Indexes", "Memos"],
        ["Harvest systems", "SelfDoc", "MDO", "Data dictionary"],
        ["Proof artifacts", "Transcripts", "Readback", "Promotion"],
    ],
    "DIAG-EVIDENCE-PIPELINE": [
        ["Source contracts", "HELP DBFs", "META/SelfDoc"],
        ["cmdhelpchk", "Runtime proofs", "manualgen"],
        ["Reader manual", "Developer manual", "Data dictionary"],
    ],
    "DIAG-TRINITY-HEADERS": [
        ["xbase.hpp", "classic DBF core"],
        ["xbase_vfp.hpp", "VFP compatibility"],
        ["xbase_64.hpp", "x64 extensions"],
    ],
    "DIAG-X64-SELF-DESCRIBING": [
        ["DBF header", "field descriptors", "terminator"],
        ["X64M metadata", "logical table name", "logical field names"],
        ["record data", "fallback tokens", "readback"],
    ],
    "DIAG-COMMAND-HARVEST": [
        ["@dottalk.usage blocks", "source comments"],
        ["raw harvest", "selected command guide"],
        ["HELP/META join", "proof status", "reader reference"],
    ],
    "DIAG-WORKAREAS-CURSOR": [
        ["work area 1", "current table", "current record"],
        ["work area 2", "alternate table", "cursor state"],
        ["RECNO", "GOTO", "SKIP", "restoration"],
    ],
    "DIAG-LOCKING-LIFECYCLE": [
        ["open table", "lock record/table", "check owner"],
        ["mutate safely", "commit/rollback", "unlock"],
    ],
    "DIAG-BUFFER-COMMIT-ROLLBACK": [
        ["TABLE BUFFER ON", "staged REPLACE", "dirty state"],
        ["COMMIT", "physical readback", "clean state"],
        ["ROLLBACK", "discard staged changes", "restore view"],
    ],
    "DIAG-INDEX-ORDER": [
        ["DBF physical order", "CNX/CDX container", "SET INDEX"],
        ["tag/order selection", "SET ORDER", "logical traversal"],
        ["LMDB backend detail", "reader sees xBase order"],
    ],
    "DIAG-MEMO-DTX": [
        ["DBF memo field", "object id/reference"],
        ["DTX sidecar", "payload blocks", "readback"],
    ],
    "DIAG-MANUAL-FAMILY": [
        ["Reader manual", "Developer handoff", "Command reference"],
        ["Anchor map", "Maturity model", "Data dictionary"],
        ["Diagram registry", "Proof transcripts", "Manualgen catalog"],
    ],
    "DIAG-MATURITY-MODEL": [
        ["M0 captured", "M1 anchored", "M2 source-supported"],
        ["M3 runtime-proven", "M4 reader-ready", "M5 reference-complete"],
        ["M6 release-reviewed"],
    ],
    "DIAG-PROOF-LOOP": [
        ["run command/script", "capture transcript"],
        ["inspect readback", "compare catalogs"],
        ["promote section", "keep proof path"],
    ],
}


SECTION_DIAGRAMS = {
    "What DotTalk++ Is": ["DIAG-ARCH-LAYERS"],
    "The Trinity Headers": ["DIAG-TRINITY-HEADERS"],
    "The x64 Workflow": ["DIAG-X64-SELF-DESCRIBING"],
    "x64 Memo Limits": ["DIAG-MEMO-DTX"],
    "Schema Rules": ["DIAG-X64-SELF-DESCRIBING"],
    "Indexes, CNX, CDX, and SET ORDER": ["DIAG-INDEX-ORDER"],
    "Work Areas and Cursor Control": ["DIAG-WORKAREAS-CURSOR"],
    "Record and Table Locking": ["DIAG-LOCKING-LIFECYCLE"],
    "Table Buffering": ["DIAG-BUFFER-COMMIT-ROLLBACK"],
    "Commit and Rollback": ["DIAG-BUFFER-COMMIT-ROLLBACK"],
    "Commands and Functions Reference": ["DIAG-COMMAND-HARVEST"],
    "Evidence Anchors and Manual Generation": ["DIAG-EVIDENCE-PIPELINE", "DIAG-MANUAL-FAMILY", "DIAG-MATURITY-MODEL", "DIAG-PROOF-LOOP"],
}


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)


def load_registry() -> dict[str, dict[str, str]]:
    rows: dict[str, dict[str, str]] = {}
    with DIAGRAM_REGISTRY.open(newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            rows[row["asset_id"]] = row
    return rows


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/segoeui.ttf",
    ]
    if bold:
        candidates = [
            "C:/Windows/Fonts/calibrib.ttf",
            "C:/Windows/Fonts/arialbd.ttf",
            "C:/Windows/Fonts/segoeuib.ttf",
        ] + candidates
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def draw_centered(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fnt, fill=(25, 35, 45)) -> None:
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 22
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        probe = f"{current} {word}".strip()
        if draw.textbbox((0, 0), probe, font=fnt)[2] <= max_width:
            current = probe
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    line_h = draw.textbbox((0, 0), "Ag", font=fnt)[3] + 4
    total_h = line_h * len(lines)
    y = y1 + ((y2 - y1 - total_h) // 2)
    for line in lines:
        tw = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text((x1 + ((x2 - x1 - tw) // 2), y), line, font=fnt, fill=fill)
        y += line_h


def create_diagram_png(asset_id: str, meta: dict[str, str]) -> Path:
    out = IMG_DIR / f"{asset_id.lower()}.png"
    rows = DIAGRAMS[asset_id]
    width, height = 1400, 820
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(38, True)
    label_font = font(24, True)
    small_font = font(21, False)
    draw.rectangle((0, 0, width - 1, height - 1), outline=(190, 202, 218), width=3)
    draw.text((60, 42), meta["purpose"], font=title_font, fill=(31, 77, 120))
    draw.text((60, 94), f'{asset_id}  |  {meta["manual_anchor"]}', font=small_font, fill=(90, 96, 104))
    top = 160
    row_h = 145 if len(rows) <= 4 else 118
    gap_y = 24
    for r, cells in enumerate(rows):
        cols = len(cells)
        margin_x = 72
        gap_x = 30
        box_w = int((width - (2 * margin_x) - (gap_x * (cols - 1))) / cols)
        y1 = top + r * (row_h + gap_y)
        for c, text in enumerate(cells):
            x1 = margin_x + c * (box_w + gap_x)
            x2 = x1 + box_w
            y2 = y1 + row_h
            fill = (232, 238, 245) if r % 2 == 0 else (244, 246, 249)
            draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=fill, outline=(135, 158, 188), width=3)
            draw_centered(draw, (x1, y1, x2, y2), text, label_font)
            if c < cols - 1:
                ax1 = x2 + 4
                ay = y1 + row_h // 2
                ax2 = x2 + gap_x - 7
                draw.line((ax1, ay, ax2, ay), fill=(72, 104, 138), width=4)
                draw.polygon([(ax2, ay), (ax2 - 12, ay - 8), (ax2 - 12, ay + 8)], fill=(72, 104, 138))
        if r < len(rows) - 1:
            cx = width // 2
            yy = y1 + row_h + 6
            draw.line((cx, yy, cx, yy + gap_y - 9), fill=(72, 104, 138), width=4)
            draw.polygon([(cx, yy + gap_y - 3), (cx - 8, yy + gap_y - 15), (cx + 8, yy + gap_y - 15)], fill=(72, 104, 138))
    img.save(out)
    return out


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = BORDER) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def table_widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(row.cells[idx])


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Right-click and update field in Word to refresh page numbers."
    sep_run = OxmlElement("w:r")
    sep_run.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(fld)
    run._r.append(instr)
    run._r.append(sep)
    paragraph._p.append(sep_run)
    run._r.append(end)


def add_update_fields_setting(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT, 14, 7),
        ("Heading 3", 12, DARK, 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    styles["Caption"].font.name = "Calibri"
    styles["Caption"].font.size = Pt(9)
    styles["Caption"].font.color.rgb = MUTED


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.text = "DotTalk++ / x64base Manual Preview V1"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = MUTED


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("DotTalk++ / x64base")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = DARK
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Manual Preview V1")
    r2.font.size = Pt(18)
    r2.font.color.rgb = ACCENT
    doc.add_paragraph("A reader-facing assembled manual with diagrams, anchor map, maturity model, command-reference seed, and generated topical index.")
    meta = doc.add_table(rows=5, cols=2)
    table_widths(meta, [1.55, 4.95])
    rows = [
        ("Status", "Inspection draft assembled from current repository manual sources."),
        ("Audience", "Human developer, AI development agent, documentation maintainer, and early reader."),
        ("Source root", str(ROOT)),
        ("Generated from", "Reader manual, anchor map, diagram registry, maturity model, command usage harvest."),
        ("Design preset", "compact_reference_guide."),
    ]
    for i, (k, v) in enumerate(rows):
        meta.cell(i, 0).text = k
        meta.cell(i, 1).text = v
        set_cell_shading(meta.cell(i, 0), LIGHT_FILL)
    doc.add_page_break()


def add_front_matter(doc: Document, registry: dict[str, dict[str, str]]) -> None:
    doc.add_heading("Table of Contents", level=1)
    add_toc(doc.add_paragraph())
    doc.add_page_break()
    doc.add_heading("Illustrations", level=1)
    p = doc.add_paragraph("The figures in this preview are anchored manual assets. They clarify current structure but do not increase a section's proof state; proof still comes from source contracts, catalogs, and runtime transcripts.")
    p.runs[0].italic = True
    table = doc.add_table(rows=1, cols=4)
    table_widths(table, [1.35, 2.35, 1.7, 1.1])
    headers = ["Figure", "Purpose", "Anchor", "Asset ID"]
    for idx, h in enumerate(headers):
        table.cell(0, idx).text = h
        set_cell_shading(table.cell(0, idx), LIGHT_FILL)
    for n, (asset_id, row) in enumerate(registry.items(), start=1):
        cells = table.add_row().cells
        cells[0].text = f"Figure {n}"
        cells[1].text = row["purpose"]
        cells[2].text = row["manual_anchor"]
        cells[3].text = asset_id
    doc.add_page_break()


def add_figure(doc: Document, asset_id: str, registry: dict[str, dict[str, str]], figure_no: int) -> int:
    row = registry[asset_id]
    path = IMG_DIR / f"{asset_id.lower()}.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.3))
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(f"Figure {figure_no}. {row['purpose']} ({asset_id}, {row['manual_anchor']}).")
    return figure_no + 1


def clean_inline(text: str) -> str:
    return re.sub(r"`([^`]+)`", r"\1", text).replace("**", "")


def add_markdown_body(doc: Document, registry: dict[str, dict[str, str]]) -> None:
    doc.add_heading("Manual Body", level=1)
    figure_no = 1
    in_code = False
    code_lines: list[str] = []
    current_h2 = ""
    skip_toc_list = False
    for raw in MANUAL.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if line.startswith("## Table of Contents"):
            skip_toc_list = True
            continue
        if skip_toc_list:
            if line.startswith("## "):
                skip_toc_list = False
            else:
                continue
        if line.startswith("```"):
            if in_code:
                if code_lines:
                    p = doc.add_paragraph()
                    run = p.add_run("\n".join(code_lines))
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(8)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            continue
        m = re.match(r"^(#{1,3})\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)), 3)
            title = clean_inline(m.group(2))
            doc.add_heading(title, level=level)
            normalized_title = re.sub(r"^\d+\.\s+", "", title)
            if level == 2:
                current_h2 = normalized_title
            for asset_id in SECTION_DIAGRAMS.get(normalized_title, []):
                figure_no = add_figure(doc, asset_id, registry, figure_no)
            continue
        bullet = re.match(r"^-\s+(.*)$", line)
        if bullet:
            doc.add_paragraph(clean_inline(bullet.group(1)), style="List Bullet")
            continue
        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            doc.add_paragraph(clean_inline(numbered.group(1)), style="List Number")
            continue
        if line.startswith("|") and line.endswith("|"):
            # Keep markdown tables readable without fragile table parsing.
            p = doc.add_paragraph()
            run = p.add_run(clean_inline(line))
            run.font.name = "Consolas"
            run.font.size = Pt(8)
            continue
        doc.add_paragraph(clean_inline(line))


def add_anchor_summary(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix A: Anchor Map Summary", level=1)
    rows = []
    for line in ANCHOR_MAP.read_text(encoding="utf-8").splitlines():
        if line.startswith("| `ANCHOR-"):
            parts = [p.strip().strip("`") for p in line.strip("|").split("|")]
            if len(parts) >= 6:
                rows.append(parts[:6])
    table = doc.add_table(rows=1, cols=4)
    table_widths(table, [1.75, 1.35, 1.65, 1.75])
    for idx, h in enumerate(["Anchor", "Layer", "State", "Manual Target"]):
        table.cell(0, idx).text = h
        set_cell_shading(table.cell(0, idx), LIGHT_FILL)
    for anchor, layer, _evidence, target, state, _next in rows:
        cells = table.add_row().cells
        cells[0].text = anchor
        cells[1].text = layer
        cells[2].text = state
        cells[3].text = target


def add_maturity_summary(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix B: Maturity Model", level=1)
    capture = False
    for line in MATURITY.read_text(encoding="utf-8").splitlines():
        if line.startswith("## Maturity Levels"):
            capture = True
            continue
        if line.startswith("## Professional Documentation Standard"):
            break
        if not capture or not line.strip() or line.startswith("|---"):
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("|"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(8)
        else:
            doc.add_paragraph(clean_inline(line))


def add_command_excerpt(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix C: Command Reference Seed", level=1)
    text = COMMAND_GUIDE.read_text(encoding="utf-8")
    intro = re.search(r"Selected command entries:.*?## Command Entries", text, re.S)
    if intro:
        for line in intro.group(0).replace("## Command Entries", "").splitlines():
            if line.strip():
                doc.add_paragraph(clean_inline(line))
    commands = re.split(r"\n---\n", text.split("## Command Entries", 1)[1])
    for block in commands[:24]:
        lines = block.strip().splitlines()
        if not lines:
            continue
        heading = lines[0].replace("### ", "").strip()
        doc.add_heading(heading, level=2)
        for line in lines[1:]:
            if not line.strip() or line.startswith("```"):
                continue
            if line.startswith("Source comments:"):
                break
            if line.startswith("Source:") or line.startswith("Owner:") or line.startswith("Category:") or line.startswith("Status:") or line.startswith("Effect:") or line.startswith("Mutates:"):
                doc.add_paragraph(clean_inline(line))
            elif line.endswith(":"):
                doc.add_heading(line.rstrip(":"), level=3)
            else:
                doc.add_paragraph(clean_inline(line))
    doc.add_paragraph("The full harvested command guide remains in DOTTALKPP_COMMAND_REFERENCE_GUIDE_V1.md.")


def add_topical_index(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Generated Topical Index", level=1)
    terms = {
        "Anchors": "Evidence Anchors and Manual Generation; Appendix A",
        "Buffering": "Table Buffering; Commit and Rollback",
        "CDX": "Indexes, CNX, CDX, and SET ORDER",
        "CNX": "Indexes, CNX, CDX, and SET ORDER",
        "Command reference": "Commands and Functions Reference; Appendix C",
        "Commit": "Commit and Rollback",
        "Cursor control": "Work Areas and Cursor Control",
        "Data dictionary": "Schema Rules; Evidence Anchors and Manual Generation",
        "DotScript": "DotScript",
        "DTX memos": "Theoretical Limits; x64 Memo Limits",
        "HELP": "Commands and Functions Reference; Evidence Anchors and Manual Generation",
        "Indexing": "Indexes, CNX, CDX, and SET ORDER",
        "Locking": "Record and Table Locking",
        "Maturity": "Evidence Anchors and Manual Generation; Appendix B",
        "MDO": "Developer Appendix; Evidence Anchors and Manual Generation",
        "META": "Commands and Functions Reference; Evidence Anchors and Manual Generation",
        "Rollback": "Commit and Rollback",
        "SelfDoc": "Evidence Anchors and Manual Generation",
        "SET ORDER": "Indexes, CNX, CDX, and SET ORDER",
        "Trinity headers": "The Trinity Headers",
        "Vectored names": "Schema Rules; The x64 Workflow",
        "Work areas": "Work Areas and Cursor Control",
        "X64M": "The x64 Workflow; Schema Rules",
    }
    table = doc.add_table(rows=1, cols=2)
    table_widths(table, [1.85, 4.65])
    table.cell(0, 0).text = "Term"
    table.cell(0, 1).text = "See"
    set_cell_shading(table.cell(0, 0), LIGHT_FILL)
    set_cell_shading(table.cell(0, 1), LIGHT_FILL)
    for term, target in sorted(terms.items(), key=lambda kv: kv[0].lower()):
        cells = table.add_row().cells
        cells[0].text = term
        cells[1].text = target


def build() -> None:
    ensure_dirs()
    registry = load_registry()
    for asset_id, row in registry.items():
        create_diagram_png(asset_id, row)
    doc = Document()
    style_doc(doc)
    add_update_fields_setting(doc)
    add_cover(doc)
    add_front_matter(doc, registry)
    add_markdown_body(doc, registry)
    add_anchor_summary(doc)
    add_maturity_summary(doc)
    add_command_excerpt(doc)
    add_topical_index(doc)
    add_footer(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
